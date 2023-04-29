from io import BytesIO
import os
from docx import Document


## 결과물로 나온 값을 word파일로 생성해주는 컴포넌트
class Converter():
    # doc.add_heading(Title, level=0)
    def __init__(self, title, content, future):
        self.title = title
        self.content = content
        self.future = future

    def setting(self):
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        doc = Document()
        doc.add_heading(self.title, level=0)
        doc.add_heading('보고 내용', level=1)
        for conten in self.content:
            doc.add_paragraph(conten)
        doc.add_heading('추후 계획', level=1)
        doc.add_paragraph(self.future)
        filename = os.path.join(desktop_path, f"{self.title}.docx")
        doc.save(filename)
